"""Build the Chinese technical report requested for the STEP assembly work."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\11049\Desktop\Model_match")
DESKTOP = Path(r"C:\Users\11049\Desktop")
OUT = DESKTOP / "STEP自动装配系统_技术工作报告.docx"

IMAGES = {
    "case1": Path(r"C:\Users\11049\Documents\xwechat_files\wxid_smn8ww007ug522_0872\temp\RWTemp\2026-07\9e20f478899dc29eb19741386f9343c8\e968d11f6c0408862f044db28d77ffb8.png"),
    "case2": Path(r"C:\Users\11049\Documents\xwechat_files\wxid_smn8ww007ug522_0872\temp\RWTemp\2026-07\9e20f478899dc29eb19741386f9343c8\e5c7a7c5ef64d70505a3a395964df434.png"),
    "case3": Path(r"C:\Users\11049\Documents\xwechat_files\wxid_smn8ww007ug522_0872\temp\RWTemp\2026-07\9e20f478899dc29eb19741386f9343c8\efe6123476c5f67e97731f5eb43cda61.png"),
    "case4": Path(r"C:\Users\11049\Documents\xwechat_files\wxid_smn8ww007ug522_0872\temp\RWTemp\2026-07\9e20f478899dc29eb19741386f9343c8\d1922d682b8d7ce10f98b694fecacf12.png"),
    "case5": ROOT / "sw" / "generalization_work" / "case5_complete_v1" / "case5_psu_hole_candidate5.png",
}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "7A5A00"


def set_font(run, size: float, *, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths: list[int], header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side in ("top", "bottom", "start", "end"):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), "80" if side in {"top", "bottom"} else "120")
                node.set(qn("w:type"), "dxa")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
    if header:
        for cell in table.rows[0].cells:
            shade(cell, LIGHT)
            for run in cell.paragraphs[0].runs:
                set_font(run, 10.5, bold=True, color=INK)


def add_para(doc: Document, text: str = "", *, bold: bool = False, color: str = "000000", size: float = 11, align=None, after: float = 6) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_font(run, size, bold=bold, color=color)


def add_bullets(doc: Document, rows: list[str]) -> None:
    for row in rows:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.10
        set_font(p.add_run(row), 11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), {1: 16, 2: 13, 3: 12}[level], bold=True, color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])


def add_figure(doc: Document, image: Path, caption: str) -> None:
    if not image.is_file():
        raise FileNotFoundError(image)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(image), width=Inches(6.25))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.keep_with_next = False
    set_font(cap.add_run(caption), 9.5, color="555555")


def set_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, DARK_BLUE),
    ):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("STEP 零件自动匹配系统 | 技术工作报告"), 8.5, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    set_font(footer.add_run("内部技术报告 | 2026-07-14"), 8.5, color="666666")


def build() -> None:
    doc = Document()
    set_style(doc)

    # Memo masthead / first page
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("技术工作报告"), 12, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run("STEP 零件自动匹配与多零件 Pose 求解"), 24, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    set_font(p.add_run("面向保守候选推荐、几何约束闭合与人工复核的阶段性交付"), 13, color="555555")

    meta = doc.add_table(rows=4, cols=2)
    values = [
        ("报告日期", "2026-07-14"),
        ("范围", "case1-case5 已知组装配姿态恢复与渲染核验"),
        ("当前交付", "case1-case4 视觉验收通过；case5 形成最优孔阵列候选，保留 review 状态"),
        ("技术原则", "降低错误自动接受；证据不足的候选交由人工复核"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        set_font(row.cells[0].paragraphs[0].add_run(label), 10.5, bold=True, color=INK)
        set_font(row.cells[1].paragraphs[0].add_run(value), 10.5)
        shade(row.cells[0], LIGHT)
    configure_table(meta, [2700, 6660], header=False)
    add_heading(doc, "结论摘要", 1)
    add_para(doc, "当前系统已形成一条以几何证据为中心的保守装配求解路线。对已知真组，case1-case4 的姿态渲染经过人工视觉确认；case5 的复杂机箱场景中，导向/边缘候选与孔阵列候选存在冲突，因此最终保留孔阵列最优候选作为人工复核结果，而非把它错误升级为自动接受。")
    add_bullets(doc, [
        "已验证能力：共轴、平面贴合、相位、键槽、插入深度、刚体附属件传播以及多视图渲染。",
        "关键改进：从“单一几何分数选 Top1”转为多证据门控；Pose 成功或无碰撞均不足以证明真实来源正确。",
        "交付策略：accepted / review / rejected / unresolved 四类输出，优先压低 false positive。",
    ])
    doc.add_page_break()

    add_heading(doc, "一、当前技术路线", 1)
    add_para(doc, "系统目标不是强行把所有零件完成分组，而是在候选召回充分的前提下，给出高可信自动接受结果与可审计的人工复核结果。当前路线以原始 STEP 的 B-Rep 几何为主，语义模型默认只生成复核说明，不参与自动裁决。")
    route = doc.add_table(rows=1, cols=3)
    for cell, text in zip(route.rows[0].cells, ["阶段", "输入与核心操作", "输出 / 边界"]):
        set_font(cell.paragraphs[0].add_run(text), 10.5, bold=True, color=INK)
    rows = [
        ("候选召回", "平面、圆柱轴、孔阵列、槽口、包围盒和局部接口特征", "优先召回；不在早期剪掉真值候选"),
        ("局部 Pose", "共轴、接触、相位、插入深度、导向面/止挡和孔距刚体关系", "产生多个可审计 Pose 假设"),
        ("组级一致性", "独立几何证据数、弱单接口、中心件结构与更大组冲突", "不足两条独立证据则进入 review"),
        ("验证与交付", "原始 B-Rep 碰撞、残差、约束闭合、多视图渲染", "accepted / review / rejected / unresolved"),
    ]
    for a, b, c in rows:
        cells = route.add_row().cells
        for cell, text in zip(cells, [a, b, c]):
            set_font(cell.paragraphs[0].add_run(text), 10)
    configure_table(route, [1500, 4540, 3320])
    add_heading(doc, "二、自动接受门控", 1)
    add_para(doc, "自动接受不是由 final_score 单独决定。当前建议门控为：collision_free、pose_status=valid、geometry_score>=0.80、独立证据数>=2、group_consistency_score>=0.70、非弱单接口、无明显全局冲突，且组规模不超过 5。任何一项缺失时，结果应进入 review 或 unresolved。", color=CAUTION)
    add_heading(doc, "三、证据与风险说明", 1)
    add_bullets(doc, [
        "孔阵列是强证据，但必须与安装面共面、孔径兼容、孔距关系和局部可装配性共同成立；仅有三孔偶然吻合不能自动接受。",
        "DeepSeek 当前作为结构化复核解释工具。若校准门未通过，它不改变候选排序或 accepted/rejected 决策。",
        "复杂厂家模型的 OCCT 精确布尔可能超时或原生崩溃；此时输出 uncertain/review，不能把未完成布尔解释为无碰撞。",
    ])

    cases = [
        ("case1：法兰与轴的基础装配", "法兰面贴合、中心轴共轴、孔阵列对称性与轴向深度共同限制姿态。该结果作为基础几何约束闭合的视觉通过样例。", "图 1  case1 多视图装配渲染（人工视觉确认）", "case1"),
        ("case2：轴、键、双法兰的相位闭合", "同时使用共轴、两法兰内侧面贴合、轴键槽与法兰键槽相位对齐。该样例说明仅靠共轴无法确定旋转相位，必须加入键槽这一独立证据。", "图 2  case2 多零件相位装配渲染（人工视觉确认）", "case2"),
        ("case3：风扇模块与载体结构", "以外壳承载结构、局部插入位置、连接器/止挡方向和无明显干涉为候选证据。该结果显示系统可处理非纯回转类的复杂外形。", "图 3  case3 模块装配渲染（人工视觉确认）", "case3"),
        ("case4：主板、CPU 与内存条", "CPU 通过插座面、边界方向和中心位置定位；内存条通过插槽方向、插入深度与局部接触定位。CPU 与内存分别验证，避免一个零件的修正破坏另一个零件的姿态。", "图 4  case4 主板关系诊断渲染（人工视觉确认）", "case4"),
        ("case5：机箱、PSU 与耳件", "case5 是当前最困难场景。PSU 的局部舱位/边缘贴合候选与孔阵列候选不完全一致。按人工选择，报告保留孔阵列最优候选：三对孔中心对应，平均残差 0.246 mm。耳件保持独立导向/止挡候选；由于尚未形成与机箱的三孔确认，整体仍为 review。", "图 5  case5 PSU 孔阵列候选渲染（最优 review 候选）", "case5"),
    ]
    for title, desc, caption, key in cases:
        doc.add_page_break()
        add_heading(doc, title, 1)
        add_para(doc, desc)
        if key == "case5":
            note = doc.add_table(rows=1, cols=1)
            shade(note.cell(0, 0), CALLOUT)
            p = note.cell(0, 0).paragraphs[0]
            set_font(p.add_run("状态说明：case5 当前是复核候选，不作为自动接受结果。孔阵列提供了有效证据，但仍需与边缘安装面、导向/止挡关系及原始 B-Rep 碰撞验证联合确认。"), 10.5, color=CAUTION)
            configure_table(note, [9360], header=False)
        add_figure(doc, IMAGES[key], caption)

    doc.add_page_break()
    add_heading(doc, "四、阶段性结果与下一步", 1)
    status = doc.add_table(rows=1, cols=3)
    for cell, text in zip(status.rows[0].cells, ["对象", "当前结果", "证据与处理建议"]):
        set_font(cell.paragraphs[0].add_run(text), 10.5, bold=True, color=INK)
    rows = [
        ("case1", "视觉通过", "保留为法兰/轴接触与共轴基准样例。"),
        ("case2", "视觉通过", "保留键槽相位与法兰面贴合的联合约束。"),
        ("case3", "视觉通过", "保留为复杂非回转模块的局部插入样例。"),
        ("case4", "视觉通过", "保留 CPU 与 DIMM 的独立插入验证，防止联动破坏。"),
        ("case5", "review", "PSU 选用孔阵列最优候选；继续识别耳件导向面、卡扣/止挡与孔阵列。"),
    ]
    for row in rows:
        cells = status.add_row().cells
        for cell, text in zip(cells, row):
            set_font(cell.paragraphs[0].add_run(text), 10)
    configure_table(status, [1500, 1800, 6060])
    add_heading(doc, "最小下一步建议", 1)
    add_bullets(doc, [
        "将 case5 的孔阵列匹配限制在由导向/止挡给出的局部 ROI 内，避免全局三孔偶然匹配。",
        "为耳件提取折弯导向面、局部槽口和端部止挡；孔阵列只做最终锁紧验证。",
        "继续维护保守分层：不确定候选进入 review，而不是为提高覆盖率强行输出完整自动装配。",
        "用真实功能装配模板补充数据集，并对 mixed-pool 场景分别统计 auto_accept_precision、false_positive_count、review_rate 与 unresolved_parts_count。",
    ])
    add_para(doc, "本报告中 case1-case4 的“通过”指人工对渲染结果的视觉确认；case5 的“最优”指当前证据下的最佳 review 候选，并不等同于已完成工程级验收。", color="555555", size=9.5, after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "STEP 自动装配系统技术工作报告"
    doc.core_properties.subject = "case1-case5 装配姿态恢复与渲染核验"
    doc.core_properties.author = "Model_match"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

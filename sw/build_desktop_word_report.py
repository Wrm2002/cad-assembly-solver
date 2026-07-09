"""Build a DOCX technical report with multi-view renders for case1-5.

This script intentionally separates two jobs:
1. CAD rendering uses the project conda environment because it has OCCT.
2. DOCX authoring can be done by the bundled document Python runtime.

When run with the project conda Python, it renders images and writes a
JSON payload for the DOCX builder. When run with a Python that has
python-docx, it also writes the DOCX.
"""

from __future__ import annotations

import json
import math
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
OUT_DIR = DESKTOP / "CAD_Assembly_Word_Report_v3"
IMG_DIR = OUT_DIR / "images"
PAYLOAD_PATH = OUT_DIR / "report_payload.json"
DOCX_PATH = OUT_DIR / "CAD装配关系识别技术路线与case求解报告.docx"

CASES = [
    ("case_1", ROOT / "sw/1", "两个法兰同轴贴合"),
    ("case_2", ROOT / "sw/2", "轴、两个法兰和键的装配"),
    ("case_3", ROOT / "sw/3", "风扇模块插入风扇笼"),
    ("case_4", ROOT / "sw/4_lightweight", "主板、内存条与 CPU 装配"),
    ("case_5", ROOT / "sw/5_lightweight", "机箱、挂耳与电源模块装配"),
]

VIEW_LABELS_ZH = {
    "iso": "等轴测",
    "xy": "正视 XY",
    "xz": "侧视 XZ",
    "yz": "俯/侧视 YZ",
    "axis_side": "case2 轴向侧视",
    "axis_end": "case2 轴向端视",
}

COLORS = [
    (70, 130, 180),
    (220, 120, 80),
    (90, 170, 110),
    (180, 140, 220),
    (210, 170, 70),
]


@dataclass
class Triangle:
    pts: list[tuple[float, float, float]]
    normal: tuple[float, float, float]
    color: tuple[int, int, int]
    component: str


def _vsub(a: tuple[float, float, float], b: tuple[float, float, float]) -> tuple[float, float, float]:
    return (a[0] - b[0], a[1] - b[1], a[2] - b[2])


def _cross(a: tuple[float, float, float], b: tuple[float, float, float]) -> tuple[float, float, float]:
    return (
        a[1] * b[2] - a[2] * b[1],
        a[2] * b[0] - a[0] * b[2],
        a[0] * b[1] - a[1] * b[0],
    )


def _dot(a: tuple[float, float, float], b: tuple[float, float, float]) -> float:
    return sum(x * y for x, y in zip(a, b))


def _norm(v: tuple[float, float, float]) -> tuple[float, float, float]:
    length = math.sqrt(sum(x * x for x in v))
    if length < 1e-12:
        return (0.0, 0.0, 1.0)
    return tuple(x / length for x in v)  # type: ignore[return-value]


def _camera_basis(view: tuple[float, float, float], up_hint: tuple[float, float, float]):
    view = _norm(view)
    if abs(_dot(view, _norm(up_hint))) > 0.96:
        up_hint = (0.0, 1.0, 0.0)
    right = _norm(_cross(view, up_hint))
    up = _norm(_cross(right, view))
    return right, up, view


def _axis_side_basis(axis: tuple[float, float, float]):
    axis = _norm(axis)
    view = _norm(_cross(axis, (0.0, 0.0, 1.0)))
    if math.sqrt(sum(x * x for x in view)) < 1e-6:
        view = _norm(_cross(axis, (0.0, 1.0, 0.0)))
    right = axis
    up = _norm(_cross(right, view))
    return right, up, view


def _transform_point(trsf: Any, point: tuple[float, float, float]) -> tuple[float, float, float]:
    from OCC.Core.gp import gp_Pnt

    p = gp_Pnt(point[0], point[1], point[2]).Transformed(trsf)
    return (p.X(), p.Y(), p.Z())


def _triangles_from_shape(shape: Any, color: tuple[int, int, int], component_name: str, *, deflection: float):
    from OCC.Core.BRep import BRep_Tool
    from OCC.Core.BRepMesh import BRepMesh_IncrementalMesh
    from OCC.Core.TopAbs import TopAbs_FACE, TopAbs_REVERSED
    from OCC.Core.TopExp import TopExp_Explorer

    try:
        from OCC.Core.TopoDS import Face as topods_Face
    except ImportError:  # pragma: no cover
        from OCC.Core.TopoDS import topods_Face

    BRepMesh_IncrementalMesh(shape, deflection, False, 0.5, True)
    triangles: list[Triangle] = []
    exp = TopExp_Explorer(shape, TopAbs_FACE)
    while exp.More():
        face = topods_Face(exp.Current())
        loc = face.Location()
        triangulation = BRep_Tool.Triangulation(face, loc)
        if triangulation:
            trsf = loc.Transformation()
            for idx in range(1, triangulation.NbTriangles() + 1):
                tri = triangulation.Triangle(idx)
                n1, n2, n3 = tri.Get()
                pts = []
                for node_index in (n1, n2, n3):
                    node = triangulation.Node(node_index)
                    pts.append(_transform_point(trsf, (node.X(), node.Y(), node.Z())))
                normal = _norm(_cross(_vsub(pts[1], pts[0]), _vsub(pts[2], pts[0])))
                if face.Orientation() == TopAbs_REVERSED:
                    normal = tuple(-x for x in normal)  # type: ignore[assignment]
                triangles.append(Triangle(pts, normal, color, component_name))
        exp.Next()
    return triangles


def _load_case_triangles(case_dir: Path, payload: dict[str, Any]) -> list[Triangle]:
    from OCC.Core.BRepBuilderAPI import BRepBuilderAPI_Transform

    sys.path.insert(0, str(ROOT / "sw"))
    from build_assembly import build_transform, load_step

    output_dir = case_dir / "known_group_output"
    triangles: list[Triangle] = []
    for index, component in enumerate(payload.get("components", [])):
        source_path = (output_dir / component["source"]).resolve()
        shape = load_step(str(source_path))
        trsf = build_transform(component.get("placement", {}))
        if trsf.Form() != 0:
            shape = BRepBuilderAPI_Transform(shape, trsf, True).Shape()
        triangles.extend(
            _triangles_from_shape(
                shape,
                COLORS[index % len(COLORS)],
                Path(component["source"]).name,
                deflection=1.5 if case_dir.name in {"3"} else 0.75,
            )
        )
    return triangles


def _project(
    triangles: list[Triangle],
    width: int,
    height: int,
    *,
    right: tuple[float, float, float],
    up: tuple[float, float, float],
    view: tuple[float, float, float],
):
    projected = []
    xs: list[float] = []
    ys: list[float] = []
    for tri in triangles:
        pts2 = []
        zs = []
        for p in tri.pts:
            x = _dot(p, right)
            y = _dot(p, up)
            z = _dot(p, view)
            pts2.append((x, y))
            zs.append(z)
            xs.append(x)
            ys.append(y)
        projected.append((tri, pts2, sum(zs) / len(zs)))
    if not xs or not ys:
        return []
    margin = 80
    min_x, max_x = min(xs), max(xs)
    min_y, max_y = min(ys), max(ys)
    scale = min(
        (width - 2 * margin) / max(max_x - min_x, 1e-9),
        (height - 2 * margin) / max(max_y - min_y, 1e-9),
    )

    def screen(pt: tuple[float, float]) -> tuple[float, float]:
        return (
            margin + (pt[0] - min_x) * scale,
            height - margin - (pt[1] - min_y) * scale,
        )

    return [(tri, [screen(pt) for pt in pts2], depth) for tri, pts2, depth in projected]


def render_image(
    triangles: list[Triangle],
    image_path: Path,
    *,
    title: str,
    right: tuple[float, float, float],
    up: tuple[float, float, float],
    view: tuple[float, float, float],
) -> None:
    from PIL import Image, ImageDraw

    width, height = 1400, 1000
    image = Image.new("RGB", (width, height), (250, 250, 248))
    draw = ImageDraw.Draw(image)
    projected = _project(triangles, width, height, right=right, up=up, view=view)
    light = _norm((0.4, -0.3, 0.85))
    for tri, pts2, _depth in sorted(projected, key=lambda row: row[2]):
        brightness = 0.45 + 0.55 * max(0.0, _dot(_norm(tri.normal), light))
        color = tuple(max(0, min(255, int(c * brightness))) for c in tri.color)
        draw.polygon(pts2, fill=color, outline=(65, 65, 65))
    draw.rectangle((20, 20, width - 20, 68), fill=(255, 255, 255), outline=(210, 210, 210))
    draw.text((34, 36), title, fill=(20, 20, 20))
    image_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(image_path)


def _case_views(case_id: str, case_dir: Path) -> dict[str, tuple[tuple[float, float, float], tuple[float, float, float], tuple[float, float, float]]]:
    views = {
        "iso": _camera_basis((1.3, -1.6, 1.0), (0.0, 0.0, 1.0)),
        "xy": _camera_basis((0.0, 0.0, 1.0), (0.0, 1.0, 0.0)),
        "xz": _camera_basis((0.0, -1.0, 0.0), (0.0, 0.0, 1.0)),
        "yz": _camera_basis((1.0, 0.0, 0.0), (0.0, 0.0, 1.0)),
    }
    if case_id == "case_2":
        pose_path = case_dir / "known_group_output" / "pose_validation.json"
        if pose_path.exists():
            pose = json.loads(pose_path.read_text(encoding="utf-8"))
            axis = None
            for item in pose.get("pose_audit", []):
                if item.get("candidate_origin") == "axial_slide_dof_search" and item.get("collision_count") == 0:
                    axis = tuple(item.get("axial_slide", {}).get("axis", []))
                    break
            if axis and len(axis) == 3:
                axis_t = (float(axis[0]), float(axis[1]), float(axis[2]))
                views["axis_side"] = _axis_side_basis(axis_t)
                views["axis_end"] = _camera_basis(axis_t, (0.0, 0.0, 1.0))
    return views


def render_all_images() -> dict[str, Any]:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    exam = json.loads((ROOT / "sw/exam_assembly_methods.json").read_text(encoding="utf-8"))
    case_by_id = {case["case_id"]: case for case in exam["cases"]}

    cases_payload = []
    for case_id, case_dir, title_zh in CASES:
        case_export = case_by_id[case_id]
        assembly_payload = json.loads((case_dir / "known_group_output" / "assembly_relations.json").read_text(encoding="utf-8"))
        triangles = _load_case_triangles(case_dir, assembly_payload)
        views = _case_views(case_id, case_dir)
        image_records = []
        for view_name, (right, up, view) in views.items():
            image_path = IMG_DIR / f"{case_id}_{view_name}.png"
            render_image(
                triangles,
                image_path,
                title=f"{case_id} | {view_name} | {case_export['diagnostics']['pose_status']} | collision={case_export['diagnostics']['collision_count']}",
                right=right,
                up=up,
                view=view,
            )
            image_records.append(
                {
                    "view": view_name,
                    "view_zh": VIEW_LABELS_ZH[view_name],
                    "path": str(image_path),
                }
            )
        cases_payload.append(
            {
                "case_id": case_id,
                "title_zh": title_zh,
                "case_dir": str(case_dir),
                "images": image_records,
                "export": case_export,
            }
        )
    payload = {
        "output_dir": str(OUT_DIR),
        "docx_path": str(DOCX_PATH),
        "exam_summary_path": str(ROOT / "sw/exam_assembly_methods.json"),
        "cases": cases_payload,
        "technical_route": [
            "输入假设：case 内零件已知属于同一装配体，任务不是混池自动分组，而是恢复直接装配边、关系标签和可验证 pose。",
            "几何特征提取：从 STEP 提取平面、圆柱面、bbox、孔/轴半径、轴线方向、局部面接触与凹槽/凸起证据。",
            "候选关系生成：两两生成 coaxial、clearance、planar_mate、planar_align、pocket_mate 候选。",
            "保守直接边选择：在已知同组前提下选择连通 skeleton，尽量避免 satellite-satellite 假边。",
            "关系标签推断：不使用文件名 token 或 case-specific override，而用几何证据组合推断中文装配方式。",
            "有界 pose 搜索：在 solver beam 基础上加入 identity pose、轴向滑移、平面内滑移和 pocket 插入深度搜索。",
            "OCCT 精确验证：用 Boolean Common 检查实体交叠；pose valid 表示物理可行，不等于自动证明工程语义完全正确。",
        ],
    }
    PAYLOAD_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    shutil.copy2(ROOT / "sw/exam_assembly_methods.json", OUT_DIR / "exam_assembly_methods.json")
    return payload


def build_docx_from_payload(payload: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CAD 装配关系识别技术路线与 case1-5 求解报告")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("输出目标：直接装配边、关系标签、可验证 pose 与多视角检查图").italic = True

    doc.add_heading("1. 当前技术路线", level=1)
    for item in payload["technical_route"]:
        p = doc.add_paragraph(style=None)
        p.style = styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.15)
        p.add_run("• ").bold = True
        p.add_run(item)

    doc.add_heading("2. 总体考试结果", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["Case", "直接边数", "验证器 Pose", "Collision", "检查 pose 数", "工程判断"]):
        hdr[i].text = text
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    def add_image_table(image_records: list[dict[str, str]]) -> None:
        def cant_split(row: Any) -> None:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))

        for idx in range(0, len(image_records), 2):
            img_table = doc.add_table(rows=0, cols=2)
            img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            img_table.style = "Table Grid"
            row = img_table.add_row().cells
            cant_split(img_table.rows[0])
            for j in range(2):
                if idx + j >= len(image_records):
                    continue
                rec = image_records[idx + j]
                p = row[j].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(rec["view_zh"]).bold = True
                p.add_run().add_break()
                p.add_run().add_picture(rec["path"], width=Inches(2.95))
            doc.add_paragraph()

    for case_index, case in enumerate(payload["cases"]):
        export = case["export"]
        diag = export["diagnostics"]
        row = table.add_row().cells
        values = [
            export["case_id"],
            str(len(export["direct_assembly_edges"])),
            diag["pose_status"],
            str(diag["collision_count"]),
            str(diag["checked_pose_count"]),
            "通过当前检查" if diag["pose_status"] == "valid" else "待复核/未自动接受",
        ]
        for i, text in enumerate(values):
            row[i].text = text
            row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph(
        "注意：这里的 pose valid 表示当前几何约束和 OCCT 碰撞检查通过；它证明“物理上不穿插且满足当前约束”，"
        "但不单独证明工程语义一定正确。case2 已按这个原则降级为待复核。"
    )

    doc.add_heading("3. 分 case 求解经过与多视角渲染", level=1)
    for case in payload["cases"]:
        export = case["export"]
        diag = export["diagnostics"]
        doc.add_heading(f"{export['case_id']}：{case['title_zh']}", level=2)
        doc.add_paragraph(
            f"零件数：{len(export['parts'])}；直接装配边：{len(export['direct_assembly_edges'])}；"
            f"pose：{diag['pose_status']}；collision：{diag['collision_count']}；"
            f"检查 pose 数：{diag['checked_pose_count']}。"
        )

        edge_lines = []
        for edge in export["direct_assembly_edges"]:
            relation_zh = " + ".join(edge.get("relation_types_zh", edge.get("relation_types", [])))
            edge_lines.append(f"{edge['parts'][0]} ↔ {edge['parts'][1]}：{relation_zh}")
        for line in edge_lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.add_run("• ").bold = True
            p.add_run(line)

        if export["case_id"] == "case_1":
            desc = "系统识别到两个法兰之间同时存在同轴圆柱面证据和平面贴合证据；identity pose 已经无碰撞，因此输出同轴贴合。"
        elif export["case_id"] == "case_2":
            desc = (
                "系统先识别 shaft 与两个 flange 的轴孔间隙配合，以及 key 与 shaft 的插槽/平面贴合。"
                "本轮确认 planar_mate 只由反平行法向生成、planar_align 只由同向法向生成，并加入了 180° 朝向候选。"
                "轴向滑移现在会对 shaft 负侧法兰生成镜像翻面候选，解决之前两个法兰主轴方向同向的问题。"
                "但最新候选中仍有一个法兰未满足有效轴向贯穿，因此最终状态保持 uncertain/review，而不是自动接受。"
            )
        elif export["case_id"] == "case_3":
            desc = (
                "系统识别风扇模块与风扇笼之间存在候选插槽/平面贴合关系。本轮修正了 planar_mate/planar_align 法向极性，"
                "并为大 STEP 增加受限的正反插入深度候选；同时将 pocket 检测限制在面积最大的结构平面子集，避免海量面聚类崩溃。"
                "最新 7 个候选仍没有同时满足真实 pocket 闭合与精确验证，因此保持 uncertain/review。"
            )
        elif export["case_id"] == "case_4":
            desc = (
                "系统识别 PCBA 与内存、CPU 之间的局部凹槽/插入和平面贴合关系；在更严格 pocket 门控后，"
                "仅凭平面贴合不足以自动接受，因此当前状态为 uncertain/review。"
            )
        else:
            desc = (
                "系统识别挂耳与机箱的同轴贴合关系，以及 PSU 与机箱的插入/平面贴合关系；在更严格 pocket 门控后，"
                "PSU 插入关系仍需要补充真实插入深度/包围关系检查，因此当前状态为 uncertain/review。"
            )
        doc.add_paragraph("求解经过：" + desc)

        images = case["images"]
        if export["case_id"] == "case_2":
            add_image_table([rec for rec in images if not rec["view"].startswith("axis_")])
            doc.add_page_break()
            doc.add_heading("case2 轴向复核视图", level=3)
            add_image_table([rec for rec in images if rec["view"].startswith("axis_")])
        else:
            add_image_table(images)
        if export["case_id"] == "case_2":
            doc.add_paragraph(
                "case2 复核结论：反装方向已经被显式建模，候选中能生成负侧法兰 180° 镜像翻面且 collision=0 的姿态；"
                "但一个法兰的轴向有效贯穿长度仍不足，因此没有升级为 valid。"
            )
        if export["case_id"] == "case_3":
            doc.add_paragraph(
                "case3 复核结论：此前“贴在笼体外侧”的 pose 已不会被自动接受。受限插入深度搜索已经执行，"
                "但尚未找到 pocket 闭合且精确验证通过的姿态，因此仍需继续修正 fan-cage 开口包围和插入深度模型。"
            )
        if case_index < len(payload["cases"]) - 1:
            doc.add_page_break()
    doc.save(DOCX_PATH)


def main() -> None:
    if "--docx-only" in sys.argv:
        payload = json.loads(PAYLOAD_PATH.read_text(encoding="utf-8"))
        build_docx_from_payload(payload)
        print(json.dumps({"output_dir": str(OUT_DIR), "docx": str(DOCX_PATH)}, ensure_ascii=False))
        return
    payload = render_all_images()
    try:
        import docx  # noqa: F401
    except Exception:
        print(json.dumps({"payload": str(PAYLOAD_PATH), "docx_skipped": True}, ensure_ascii=False))
        return
    build_docx_from_payload(payload)
    print(json.dumps({"output_dir": str(OUT_DIR), "docx": str(DOCX_PATH)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
